import os
from docx.shared import Pt
from docx import Document


def opendocx():
    run = doc.add_paragraph().add_run("Just a word.") 
    run.font.size = Pt(24)
    run.bold = False
    for i in sub.keys(): 
        doc.save(f"{sub[i]}7.docx")
        doc.save(f"{sub[i]}8.docx") 
        doc.save(f"{sub[i]}9.docx") 
        

path = os.path.abspath(__file__)
path = path.replace("\generator.py", "")


folder = list(os.walk(path))
print(folder)

doc = Document()

sub = {
    "Алгебра": "Algebra",
    "Физика": "Physics",
    "Русский": "Russian",
    "Биология": "Biology",
    "География": "Geography",
    "Литература": "Literature",
    "Геометрия": "Geometry",
    "История": "History",
    "Обществознание": "Social_Studies",
    "ОБЖ": "OBJ",
    "Информатика": "Information_science",
    "Программирование": "Programming",
}
opendocx()
