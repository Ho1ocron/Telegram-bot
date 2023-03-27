import os
from docx.shared import Pt
from docx import Document


def opendocx():
    run = doc.add_paragraph().add_run("Just a word.") 
    run.font.size = Pt(24)
    run.bold = False
    for i in sub.keys(): 
        doc.save(f"{sub[i]}.docx") 
        

path = os.path.abspath(__file__)
path = path.replace("\generator.py", "")


folder = list(os.walk(path))
subj = [i for i in folder[0][2]]

doc = Document()

sub = {
    "Алгебра": "Algebra",
    "Физика": "Physics",
    "Русский": "Russian",
    "Биология": "Biology",
    "География": "Geography",
    "Литература": "Literature",
    "Геометрия": "Geometry",
    "История": "History",
    "Обществознание": "Social_Studies",
    "ОБЖ": "OBJ",
    "Информатика": "IT",
    "Программирование": "Programming"
}

if __name__ == "__main__":
    for i in sub.values():
        i = f"{i}.docx"
        if i in subj:
            break
        else:
            opendocx()