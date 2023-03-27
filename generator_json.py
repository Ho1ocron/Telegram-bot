import json
import os
from docx.shared import Pt
from docx import Document


def opendocx():
    run = doc.add_paragraph().add_run("Just a word.") 
    run.font.size = Pt(24)
    run.bold = False
    for i in sub.keys(): 
        doc.save(f"{sub[i]}.docx") 


sub = open("test.json")
sub = json.load(sub)

path = os.path.abspath(__file__)
path = path.replace("\generator_json.py", "")
print(path)

folder = list(os.walk(path))
subj = [s for s in folder[0][2]]

doc = Document()

if __name__ == "__main__":
    for i in sub.values():
        i = f"{i}.docx"
        if i in subj:
            break
        else:
            opendocx()